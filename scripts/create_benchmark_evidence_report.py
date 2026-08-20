"""Build the reproducible Kernelyra benchmark evidence report.

The report intentionally distinguishes measured wins from unmeasured frameworks.
It is a release-quality evidence document, not a marketing comparison sheet.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "reports" / "Kernelyra_benchmark_evidence_2026-08-20.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
CAUTION = "7A5A00"
RISK = "9B1C1C"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def set_run_font(run, *, size: float = 11, color: str = "000000", bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, *, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table_pr = table._tbl.tblPr
    table_width = table_pr.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_pr.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")
    table_indent = table_pr.first_child_found_in("w:tblInd")
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_pr.append(table_indent)
    table_indent.set(qn("w:w"), str(TABLE_INDENT))
    table_indent.set(qn("w:type"), "dxa")
    layout = table_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_width = cell._tc.tcPr.tcW
            tc_width.set(qn("w:w"), str(widths[index]))
            tc_width.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    tr_pr.append(element)


def border_table(table) -> None:
    table_pr = table._tbl.tblPr
    borders = table_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "C9D2DE")


def write_cell(cell, text: str, *, bold: bool = False, color: str = "000000", size: float = 9.2) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    border_table(table)
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        set_cell_shading(cell, LIGHT_GRAY)
        write_cell(cell, header, bold=True, color=INK, size=9.2)
    repeat_header(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values, strict=True):
            write_cell(cell, value)


def add_paragraph(doc: Document, text: str = "", *, bold_label: str | None = None, color: str = "000000", italic: bool = False, size: float = 11) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    if bold_label:
        label = paragraph.add_run(bold_label)
        set_run_font(label, size=size, color=color, bold=True)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, italic=italic)


def add_source(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_run_font(run, size=8.5, color=MUTED, italic=True)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    style = doc.styles[f"Heading {level}"]
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, size={1: 16, 2: 13, 3: 12}[level], color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level], bold=True)


def add_callout(doc: Document, heading: str, text: str, *, color: str = INK) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH])
    border_table(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.1
    first = paragraph.add_run(f"{heading} ")
    set_run_font(first, size=10.5, color=color, bold=True)
    second = paragraph.add_run(text)
    set_run_font(second, size=10.5, color="000000")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_page_number(paragraph) -> None:
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def set_page_furniture(document: Document) -> None:
    section = document.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("KERNELYRA  |  TECHNICAL EVIDENCE REPORT")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    first = footer.add_run("Internal benchmark evidence  |  Page ")
    set_run_font(first, size=8.5, color=MUTED)
    add_page_number(footer)


def configure_styles(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for level, size, before, after, color in (
        (1, 16, 16, 8, BLUE),
        (2, 13, 12, 6, BLUE),
        (3, 12, 8, 4, DARK_BLUE),
    ):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def build() -> Path:
    document = Document()
    configure_styles(document)
    set_page_furniture(document)
    properties = document.core_properties
    properties.title = "Kernelyra Benchmark Evidence Report"
    properties.subject = "Measured performance evidence, scope and comparison gates"
    properties.author = "Kernelyra benchmark harness"
    properties.comments = "Generated from reproducible local benchmark scripts."

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Kernelyra: доказательства производительности")
    set_run_font(run, size=24, color=INK, bold=True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run("Технический отчёт о фактических измерениях, ограничениях и честном сравнении с ML-стеками")
    set_run_font(run, size=12.5, color=MUTED)
    add_paragraph(document, "Дата: 20 августа 2026  |  Платформа: Windows 11, CPU-only, 12 логических потоков, 7.7 GiB RAM, без обнаруженного GPU.", size=9.5, color=MUTED)

    add_callout(
        document,
        "Главный вывод.",
        "У Kernelyra уже есть измеримые выигрыши в отдельных нативных путях, но текущие данные НЕ доказывают, что библиотека лучше JAX, TensorFlow, PyTorch или всех других стеков. В прямом тесте против NumPy получено 1.94x, а не целевые 2x; преимущество по дополнительной рабочей памяти не подтверждено.",
        color=RISK,
    )

    add_heading(document, "1. Что именно можно утверждать сейчас")
    add_paragraph(
        document,
        "Отчёт отделяет два разных типа фактов: (1) прямое сравнение одного алгоритма с одной реализацией на одинаковых данных; (2) внутренние оптимизации Kernelyra. Только первый тип позволяет говорить о превосходстве над конкретной библиотекой.",
    )
    add_table(
        document,
        ["Проверка", "Kernelyra", "База сравнения", "Измеренный итог", "Допустимый вывод"],
        [
            [
                "Бинарная логистическая регрессия; 4096 x 2048 float32; 40 full-batch шагов; 5 независимых прогонов.",
                "612.09 обновлений/с; accuracy 0.60864; extra working set 126,976 B.",
                "NumPy vectorized: 316.30 обновлений/с; accuracy 0.61255; extra working set 81,920 B.",
                "1.935x быстрее; accuracy -0.391 п.п.; дополнительная память выше на 45,056 B.",
                "Есть ускорение CPU-ядра для этой задачи. Нельзя заявлять 2x или экономию RAM.",
            ],
            [
                "Bulk random updates; 4096 x 28; batch 64; 5000 шагов; 5 прогонов.",
                "Bulk API: 301,462.70 обновлений/с; accuracy 0.99219.",
                "Тот же Kernelyra, вызов по одному шагу: 109,465.00 обновлений/с; accuracy 0.99219.",
                "2.754x; разница accuracy 0.000 п.п.",
                "Подтверждена внутренняя оптимизация dispatch, не победа над внешней библиотекой.",
            ],
            [
                "Multiclass native kernel; 4096 x 512; 3 класса; 30 шагов; 5 прогонов.",
                "OpenMP (8 threads): 585.35 обновлений/с; accuracy 0.39331.",
                "Тот же native kernel, 1 thread: 188.69 обновлений/с; accuracy 0.39331.",
                "3.102x; accuracy совпадает.",
                "Подтверждён выигрыш от OpenMP в этом ядре; не сравнение с JAX/TF/PyTorch.",
            ],
            [
                "Hugging Face parquet I/O + policy; один доступный shard 345,319,824 B.",
                "Rust variable policy: 979.56 MiB/с.",
                "Python fixed 4 MiB chunks: 1,027.32 MiB/с.",
                "0.954x от baseline.",
                "В этом I/O тесте выигрыша нет; оптимизацию нельзя считать доказанной.",
            ],
        ],
        [1780, 1840, 1840, 1800, 2100],
    )
    add_source(
        document,
        "Воспроизводимые источники: scripts/benchmark_native_vs_numpy.py; scripts/benchmark_native_bulk_steps.py; scripts/benchmark_native_multiclass.py; .benchmarks/hf-higgs-5g/partial-report.json. Команды и значения приведены без округления до вывода в таблице.",
    )

    add_heading(document, "2. Проверка жизнеспособности на HIGGS")
    add_paragraph(
        document,
        "Это не benchmark против другой библиотеки, а проверка того, что нативный путь обучается, сохраняет checkpoint и проходит hold-out оценку. CSV содержит 98,050 записей (без заголовка). Запуск выполнен как binary classification, native linear backend, float32, лимиты CPU 80% и RAM 50%, 8,000 шагов.",
    )
    add_table(
        document,
        ["Показатель", "Значение", "Интерпретация"],
        [
            ["Статус", "completed; termination_reason = max_steps", "Без аварийной деградации; предел шагов достигнут раньше целевой метрики."],
            ["Best validation score", "0.63652 на шаге 7000", "Лучшая сохранённая checkpoint-модель."],
            ["Hold-out accuracy", "0.62911", "Качество на test split конкретного небольшого HIGGS CSV."],
            ["Hold-out ROC-AUC", "0.67389", "Есть работоспособная вероятностная модель; это не SOTA-сравнение."],
            ["Model Guard", "stable; degradation_streak = 0", "Защитная политика не обнаружила серии ухудшений."],
        ],
        [2050, 1900, 5410],
    )
    add_source(
        document,
        "Источник: .benchmarks/flat-higgs/workspace-avx2/.kernelyra/runs.sqlite3, run 334e72b183; checkpoint 334e72b183.npz.json. Dataset: .benchmarks/flat-higgs/higgs.csv.",
    )

    add_heading(document, "3. Статус сравнения с внешними библиотеками")
    add_paragraph(
        document,
        "Фраза «лучше всех библиотек» возможна только после симметричных измерений: одинаковая задача, модель, dtype, seed, data split, warm-up, версия пакета, устройство и критерий качества. Ниже - текущая матрица доказательств, а не прогноз.")
    add_table(
        document,
        ["Стек", "Прямой замер в этом отчёте", "Что требуется для корректного сравнения", "Статус вывода"],
        [
            ["NumPy", "Да: vectorized logistic regression на CPU.", "Повторить не менее 10 раз; зафиксировать BLAS/CPU affinity; измерять пик RSS отдельно от рабочих массивов.", "Kernelyra быстрее в одном узком тесте (1.94x), но не 2x и не лучше по памяти."],
            ["JAX", "Нет.", "Отдельно cold compile, warm JIT runtime, device transfer и block_until_ready(); float32 и одна целевая функция.", "Никакого заявления о превосходстве пока делать нельзя."],
            ["TensorFlow / Keras", "Нет.", "Eager и tf.function/XLA режимы; tf.data pipeline; одинаковая архитектура, precision и GPU memory policy.", "Никакого заявления о превосходстве пока делать нельзя."],
            ["PyTorch", "Нет.", "Eager и torch.compile; DataLoader warm-up; CPU OpenMP/GPU CUDA events; одинаковые optimizer и batch.", "Никакого заявления о превосходстве пока делать нельзя."],
            ["scikit-learn", "Нет.", "Сопоставлять только идентичные линейные/SGD модели и критерии качества; это CPU baseline, не общий DL-стек.", "Не измерено."],
            ["XGBoost / LightGBM", "Нет и не является прямым аналогом.", "Только на одинаковой табличной задаче, но деревья и линейные модели решают задачу по-разному.", "Сравнение возможно по outcome, но не по одному step/s."],
        ],
        [1360, 1800, 3350, 2850],
    )
    add_source(
        document,
        "Методологические источники: JAX Benchmarking (https://docs.jax.dev/en/latest/benchmarking.html); TensorFlow Profiler guide (https://www.tensorflow.org/guide/profiler); PyTorch Performance Tuning Guide (https://docs.pytorch.org/tutorials/recipes/recipes/tuning_guide.html); NumPy performant code guide (https://numpy.org/doc/stable/user/basics.performant_code.html). Проверено 20.08.2026.",
    )

    add_heading(document, "4. Почему универсальное сравнение сейчас недействительно")
    add_paragraph(document, "JAX может выиграть на скомпилированных GPU-вычислениях, но имеет отдельную стоимость JIT-компиляции и асинхронный dispatch; измерение без block_until_ready недостоверно. TensorFlow и PyTorch меняют результат в зависимости от графовой компиляции, data pipeline, CUDA kernels и mixed precision. NumPy использует векторизацию и может опираться на разные BLAS/SIMD-сборки. Поэтому цифра из CPU logistic-regression benchmark не переносится автоматически на LLM, CNN, GPU или distributed training.")
    add_callout(
        document,
        "Запрещённая формулировка до завершения матрицы.",
        "«Kernelyra быстрее JAX, TensorFlow и PyTorch во всех задачах». Текущие данные этого не поддерживают. Корректная формулировка: «Kernelyra показала 1.94x ускорение относительно NumPy в воспроизводимом CPU microbenchmark и отдельные внутренние ускорения до 3.10x».",
        color=CAUTION,
    )

    add_heading(document, "5. Ворота для будущего заявления «лучше»")
    add_table(
        document,
        ["Гейт", "Критерий", "Текущий статус"],
        [
            ["Скорость", "Медиана >= 2.0x в 10 прогонов на одной задаче; отдельно считать compile, transfer и training runtime.", "Не пройден для NumPy: 1.935x в текущем прогоне."],
            ["Качество", "Разница primary metric не хуже заранее заданного допуска; evaluation only on the same held-out split.", "Пройден для CPU NumPy microbenchmark: -0.391 п.п. в пределах исходного допуска 1 п.п."],
            ["Память", "Пиковый RSS/VRAM ниже или равен, измерен внешним профилировщиком с одинаковыми входными данными.", "Не пройден для CPU NumPy microbenchmark: extra working set выше."],
            ["Репродуцируемость", "Зафиксированные версии, seed, hardware manifest, 10+ независимых прогонов, опубликованные raw JSON результаты.", "Частично: есть скрипты, seed и hardware manifest; raw matrix для внешних стеков отсутствует."],
            ["Область применения", "Отдельные результаты для tabular CPU, tabular GPU, vision, transformer/LLM и streaming I/O.", "Не пройден: измерены только узкие tabular/streaming сценарии."],
        ],
        [1800, 4750, 2810],
    )

    add_heading(document, "6. Следующий воспроизводимый эксперимент")
    add_paragraph(document, "Первым нужно закрыть не «все библиотеки», а одну честную матрицу: CPU logistic regression и MLP, затем GPU MLP. Для каждого стека фиксируются версии, same float32 data, same parameter count, same optimizer, 10 warm runs и 10 measured runs. Для JAX измеряются cold compile и warm runtime раздельно; для GPU перед записью времени вызывается синхронизация; для TensorFlow и PyTorch отдельно публикуются eager и compiled результаты. Решение о заявлении принимается только после публикации JSON-логов и quality gates.")

    add_heading(document, "7. Воспроизведение текущих чисел")
    add_table(
        document,
        ["Цель", "Команда"],
        [
            ["NumPy vs native", "PYTHONPATH=src python scripts/benchmark_native_vs_numpy.py --runs 5"],
            ["Bulk API", "PYTHONPATH=src python scripts/benchmark_native_bulk_steps.py --runs 5"],
            ["OpenMP multiclass", "PYTHONPATH=src python scripts/benchmark_native_multiclass.py --runs 5"],
            ["HIGGS run", "PYTHONPATH=src python examples/guarded_higgs_training.py .benchmarks/flat-higgs/higgs.csv --workspace <new-workspace>"],
        ],
        [2200, 7160],
    )
    add_source(document, "Примечание: у PowerShell задайте переменную окружения так: $env:PYTHONPATH='src'. Для источника HIGGS результат зависит от доступного CSV; входной файл из текущего workspace не включён в итоговый документ.")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
